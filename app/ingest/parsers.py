from __future__ import annotations

import hashlib
import re
from pathlib import Path
from typing import Iterable

import requests
import trafilatura
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from ebooklib import ITEM_DOCUMENT, epub
from pypdf import PdfReader

from app.models import Document


def _stable_id(value: str) -> str:
    return hashlib.sha256(value.encode("utf-8")).hexdigest()[:16]


def _normalize_text(text: str) -> str:
    text = text.replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()


def _guess_topic(path: Path) -> str | None:
    parent = path.parent.name.strip()
    return parent or None


def parse_txt(path: Path) -> Document:
    text = path.read_text(encoding="utf-8", errors="ignore")
    doc_id = _stable_id(f"txt:{path.as_posix()}")
    return Document(
        doc_id=doc_id,
        text=_normalize_text(text),
        source=path.as_posix(),
        title=path.stem,
        topic=_guess_topic(path),
        metadata={"format": "txt"},
    )


def parse_pdf(path: Path) -> Document:
    reader = PdfReader(path)
    pages: list[str] = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    doc_id = _stable_id(f"pdf:{path.as_posix()}")
    return Document(
        doc_id=doc_id,
        text=_normalize_text("\n\n".join(pages)),
        source=path.as_posix(),
        title=path.stem,
        topic=_guess_topic(path),
        metadata={"format": "pdf", "page_count": len(reader.pages)},
    )


def parse_docx(path: Path) -> Document:
    doc = DocxDocument(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    doc_id = _stable_id(f"docx:{path.as_posix()}")
    return Document(
        doc_id=doc_id,
        text=_normalize_text("\n".join(paragraphs)),
        source=path.as_posix(),
        title=path.stem,
        topic=_guess_topic(path),
        metadata={"format": "docx"},
    )


def parse_epub(path: Path) -> Document:
    book = epub.read_epub(str(path))
    text_blocks: list[str] = []
    for item in book.get_items_of_type(ITEM_DOCUMENT):
        soup = BeautifulSoup(item.get_content(), "html.parser")
        text_blocks.append(soup.get_text(separator="\n"))

    title = None
    if book.title:
        title = str(book.title)
    doc_id = _stable_id(f"epub:{path.as_posix()}")
    return Document(
        doc_id=doc_id,
        text=_normalize_text("\n\n".join(text_blocks)),
        source=path.as_posix(),
        title=title or path.stem,
        topic=_guess_topic(path),
        metadata={"format": "epub"},
    )


def parse_html_file(path: Path) -> Document:
    html = path.read_text(encoding="utf-8", errors="ignore")
    soup = BeautifulSoup(html, "html.parser")
    text = soup.get_text(separator="\n")
    title = soup.title.string.strip() if soup.title and soup.title.string else path.stem
    doc_id = _stable_id(f"html:{path.as_posix()}")
    return Document(
        doc_id=doc_id,
        text=_normalize_text(text),
        source=path.as_posix(),
        title=title,
        topic=_guess_topic(path),
        metadata={"format": "html"},
    )


def parse_web_article(url: str, timeout_s: float = 10.0) -> Document:
    downloaded = trafilatura.fetch_url(url)
    extracted = trafilatura.extract(downloaded) if downloaded else None
    if not extracted:
        response = requests.get(url, timeout=timeout_s)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, "html.parser")
        extracted = soup.get_text(separator="\n")
    doc_id = _stable_id(f"url:{url}")
    return Document(
        doc_id=doc_id,
        text=_normalize_text(extracted),
        source=url,
        title=url,
        metadata={"format": "web"},
    )


def parse_path(path: Path) -> Document:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return parse_pdf(path)
    if suffix == ".docx":
        return parse_docx(path)
    if suffix == ".epub":
        return parse_epub(path)
    if suffix in {".txt", ".md"}:
        return parse_txt(path)
    if suffix in {".html", ".htm"}:
        return parse_html_file(path)
    raise ValueError(f"Unsupported file extension: {path.suffix}")


def iter_supported_files(root: Path) -> Iterable[Path]:
    allowed = {".pdf", ".docx", ".epub", ".txt", ".md", ".html", ".htm"}
    for path in root.rglob("*"):
        if path.is_file() and path.suffix.lower() in allowed:
            yield path
